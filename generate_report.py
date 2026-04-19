"""Генерация отчёта по лабораторной работе в формате .docx."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_PATH = r"C:\Users\sbogo\Databases_App\Отчёт_по_лабораторной.docx"


def set_default_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(13)
    return h


def add_para(doc: Document, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = bold
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_screenshot_note(doc: Document, caption: str, what_to_capture: str):
    """Пометка для вставки скриншота."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"[ МЕСТО ДЛЯ СКРИНШОТА: {caption} ]")
    run.bold = True
    run.italic = True
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = 1.5
    run2 = p2.add_run(f"Что снять: {what_to_capture}")
    run2.italic = True
    run2.font.size = Pt(12)
    run2.font.name = "Times New Roman"
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # подпись рисунка
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.line_spacing = 1.5
    run3 = p3.add_run(f"Рисунок — {caption}")
    run3.italic = True
    run3.font.size = Pt(13)
    run3.font.name = "Times New Roman"


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    for ridx, row in enumerate(rows, start=1):
        for cidx, val in enumerate(row):
            cell = t.rows[ridx].cells[cidx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
    return t


def build():
    doc = Document()
    set_default_font(doc)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    # ── Цель работы ──────────────────────────────────────────────────────────
    add_heading(doc, "1. Цель работы", level=1)
    add_para(
        doc,
        "Цель лабораторной работы — спроектировать и реализовать клиент-серверное "
        "приложение для работы с реляционной базой данных, поддерживающее базовые "
        "CRUD-операции над всеми сущностями предметной области, а также "
        "многомерный (OLAP) анализ данных с возможностью построения аналитических "
        "отчётов, визуализации результатов и экспорта выборок во внешние форматы.",
    )

    # ── Постановка задачи ────────────────────────────────────────────────────
    add_heading(doc, "2. Постановка задачи", level=1)
    add_para(doc, "В рамках работы необходимо решить следующие задачи:")
    add_bullet(doc, "Разработать схему реляционной базы данных, отражающую выбранную предметную область (документооборот и бюджетирование организации).")
    add_bullet(doc, "Реализовать серверную часть, предоставляющую REST API для работы с данными.")
    add_bullet(doc, "Реализовать клиентскую часть — одностраничное веб-приложение с удобным интерфейсом для конечного пользователя.")
    add_bullet(doc, "Обеспечить полный набор CRUD-операций (создание, чтение, изменение, удаление) для всех основных сущностей.")
    add_bullet(doc, "Реализовать OLAP-операции: roll-up, slice, dice, drill-down, cross-tabulation.")
    add_bullet(doc, "Добавить визуализацию аналитических данных (диаграммы).")
    add_bullet(doc, "Предусмотреть возможность экспорта отчётов во внешние форматы (CSV, XLSX).")

    # ── Предметная область ───────────────────────────────────────────────────
    add_heading(doc, "3. Описание предметной области", level=1)
    add_para(
        doc,
        "Предметная область — автоматизация учёта документооборота и бюджетирования "
        "в организации. Система позволяет вести учёт сотрудников и отделов, фиксировать "
        "бюджетные планы и фактические расходы в разрезе статей и кварталов, регистрировать "
        "документы, связанные с контрагентами, а также формировать аналитические отчёты "
        "для принятия управленческих решений.",
    )
    add_para(doc, "В базе данных выделены следующие основные сущности:")
    add_bullet(doc, "«Должность» (position) — справочник должностей с указанием грейда и минимального оклада.")
    add_bullet(doc, "«Отдел» (department) — структурное подразделение организации с указанием руководителя.")
    add_bullet(doc, "«Сотрудник» (employee) — работники организации, привязанные к отделу и должности.")
    add_bullet(doc, "«Статья бюджета» (budget_item) — справочник статей расходов с категорией.")
    add_bullet(doc, "«Тип документа» (doc_type) — справочник типов документов, срок хранения и признак обязательного согласования.")
    add_bullet(doc, "«Контрагент» (contractor) — внешние организации-партнёры, идентифицируемые по ИНН.")
    add_bullet(doc, "«Бюджет» (budget) — плановые и фактические суммы по отделу, статье, году и кварталу.")
    add_bullet(doc, "«Документ» (document) — первичные документы с привязкой к отделу, статье, типу, контрагенту и ответственному сотруднику.")

    # ── Схема БД ─────────────────────────────────────────────────────────────
    add_heading(doc, "4. Схема базы данных", level=1)
    add_para(
        doc,
        "База данных реализована в СУБД PostgreSQL и включает 8 таблиц, связанных "
        "отношениями «один-ко-многим» через внешние ключи. Таблица department "
        "содержит внешний ключ на employee (руководитель отдела), что формирует "
        "циклическую зависимость, разрешаемую порядком вставки данных.",
    )
    add_screenshot_note(
        doc,
        caption="ER-диаграмма базы данных",
        what_to_capture="Откройте pgAdmin → ПКМ по БД Laba3 → «ERD for database» (или вкладка ERD Tool) и сделайте скриншот диаграммы со всеми таблицами и связями. Альтернативно — постройте диаграмму в dbdiagram.io.",
    )
    add_para(doc, "Соответствие сущностей и таблиц базы данных приведено в таблице 1.")

    add_table(
        doc,
        headers=["Таблица", "Назначение", "Ключевые поля"],
        rows=[
            ["position", "Справочник должностей", "position_id (PK)"],
            ["doc_type", "Справочник типов документов", "type_id (PK)"],
            ["budget_item", "Справочник статей бюджета", "item_id (PK)"],
            ["contractor", "Справочник контрагентов", "contr_inn (PK)"],
            ["department", "Отделы организации", "dept_id (PK), head_emp_id (FK)"],
            ["employee", "Сотрудники", "emp_id (PK), dept_id (FK), position_id (FK)"],
            ["budget", "План и факт по бюджету", "(dept_id, item_id, budget_year, budget_quarter) (PK)"],
            ["document", "Первичные документы", "doc_id (PK), dept_id, item_id, type_id, contr_inn, resp_emp_id (FK)"],
        ],
    )

    add_screenshot_note(
        doc,
        caption="Структура таблиц в pgAdmin",
        what_to_capture="Раскройте в pgAdmin дерево БД Laba3 → Schemas → public → Tables и сделайте скриншот, где видны все 8 таблиц. Дополнительно можно снять структуру одной-двух ключевых таблиц (employee, budget) с их столбцами.",
    )

    # ── Используемые технологии ─────────────────────────────────────────────
    add_heading(doc, "5. Используемые технологии", level=1)
    add_para(doc, "При реализации приложения использованы следующие технологии:")
    add_bullet(doc, "СУБД: PostgreSQL 16 — хранение данных, поддержка транзакций и сложных аналитических запросов.")
    add_bullet(doc, "Язык серверной части: Python 3.12.")
    add_bullet(doc, "Фреймворк REST API: FastAPI — декларативное описание маршрутов, автоматическая OpenAPI-документация.")
    add_bullet(doc, "Драйвер БД: asyncpg — асинхронный доступ к PostgreSQL с пулом соединений.")
    add_bullet(doc, "Валидация данных: Pydantic v2 — модели запросов/ответов.")
    add_bullet(doc, "Язык клиентской части: JavaScript.")
    add_bullet(doc, "Фреймворк SPA: Vue 3 с Composition API.")
    add_bullet(doc, "Маршрутизация: Vue Router.")
    add_bullet(doc, "Стилизация: Tailwind CSS.")
    add_bullet(doc, "Сборщик: Vite.")
    add_bullet(doc, "Экспорт XLSX: библиотека openpyxl.")

    # ── Архитектура ─────────────────────────────────────────────────────────
    add_heading(doc, "6. Архитектура приложения", level=1)
    add_para(
        doc,
        "Приложение построено по трёхзвенной архитектуре «клиент — сервер — база "
        "данных». Клиент и сервер взаимодействуют по HTTP через REST API, "
        "обмениваясь данными в формате JSON. Сервер соединяется с PostgreSQL через "
        "пул асинхронных соединений asyncpg. Такая схема обеспечивает разделение "
        "представления, бизнес-логики и хранения данных.",
    )
    add_screenshot_note(
        doc,
        caption="Схема архитектуры приложения",
        what_to_capture="Нарисуйте блок-схему из трёх блоков: «Браузер (Vue SPA)» → «FastAPI-сервер» → «PostgreSQL». Можно сделать в draw.io, Miro или PowerPoint и вставить сюда.",
    )
    add_para(
        doc,
        "Серверная часть разделена на модули-роутеры по каждой сущности "
        "(position, department, employee, budget, document, contractor, "
        "doc_type, budget_item), а также отдельные модули olap (аналитические "
        "запросы) и export (выгрузка отчётов в CSV/XLSX).",
    )
    add_para(
        doc,
        "Клиентская часть состоит из страниц-представлений для каждой сущности "
        "(Departments, Employees, Documents, Budget, Contractors, Positions), "
        "дашборда и страницы OLAP-анализа. Общие компоненты вынесены: "
        "DataTable (универсальная таблица), FormModal (модальная форма), "
        "Sidebar (боковое меню), OlapChart, OlapTable, OlapControls.",
    )

    # ── Запуск приложения ───────────────────────────────────────────────────
    add_heading(doc, "7. Запуск приложения", level=1)
    add_para(doc, "Серверная часть запускается командой uvicorn main:app --reload из каталога backend, клиентская — npm run dev из каталога frontend. Сервер поднимается на порту 8000, клиент — на 5173.")
    add_screenshot_note(
        doc,
        caption="Запуск backend-сервера",
        what_to_capture="Терминал с запущенным uvicorn, где видны строки «Application startup complete» и «Uvicorn running on http://127.0.0.1:8000».",
    )
    add_screenshot_note(
        doc,
        caption="Запуск frontend-сервера",
        what_to_capture="Терминал с запущенным Vite, где видны строки «VITE v… ready in …» и адрес http://localhost:5173.",
    )
    add_screenshot_note(
        doc,
        caption="Swagger UI — автоматическая документация API",
        what_to_capture="Откройте в браузере http://127.0.0.1:8000/docs и сделайте скриншот списка эндпоинтов (видны разделы position, department, employee, budget, document, olap, export и т. д.).",
    )

    # ── Интерфейс и CRUD ─────────────────────────────────────────────────────
    add_heading(doc, "8. Пользовательский интерфейс и CRUD-операции", level=1)

    add_heading(doc, "8.1. Главная страница (Dashboard)", level=2)
    add_para(
        doc,
        "Главная страница содержит сводную информацию по организации: количество "
        "сотрудников, отделов, документов, контрагентов, а также ссылки на "
        "основные разделы приложения.",
    )
    add_screenshot_note(
        doc,
        caption="Главная страница приложения",
        what_to_capture="Откройте http://localhost:5173/ и сделайте скриншот дашборда с плитками счётчиков и боковым меню.",
    )

    add_heading(doc, "8.2. Раздел «Отделы»", level=2)
    add_para(
        doc,
        "Раздел «Отделы» предоставляет полный набор CRUD-операций: просмотр списка "
        "отделов с пагинацией, добавление нового отдела, редактирование и удаление "
        "существующих. Для каждого отдела отображается название, количество "
        "сотрудников, этаж, телефон и руководитель.",
    )
    add_screenshot_note(
        doc,
        caption="Список отделов",
        what_to_capture="Перейдите в раздел «Отделы», сделайте скриншот таблицы со всеми столбцами и кнопками управления (Добавить / Редактировать / Удалить).",
    )
    add_screenshot_note(
        doc,
        caption="Форма добавления/редактирования отдела",
        what_to_capture="Нажмите «Добавить» или «Редактировать» у любого отдела и сделайте скриншот модальной формы со всеми полями.",
    )

    add_heading(doc, "8.3. Раздел «Сотрудники»", level=2)
    add_para(
        doc,
        "Раздел «Сотрудники» позволяет вести учёт работников с указанием ФИО, "
        "отдела, должности, оклада, даты приёма, контактных данных. Реализованы "
        "поиск и сортировка по столбцам.",
    )
    add_screenshot_note(
        doc,
        caption="Список сотрудников",
        what_to_capture="Скриншот страницы «Сотрудники» с таблицей (желательно с применённым поиском или сортировкой).",
    )
    add_screenshot_note(
        doc,
        caption="Форма добавления сотрудника",
        what_to_capture="Скриншот модальной формы создания сотрудника со всеми полями (ФИО, отдел, должность, оклад и т. д.).",
    )

    add_heading(doc, "8.4. Раздел «Должности»", level=2)
    add_screenshot_note(
        doc,
        caption="Справочник должностей",
        what_to_capture="Скриншот страницы «Должности» с таблицей должностей (ID, название, грейд, минимальный оклад).",
    )

    add_heading(doc, "8.5. Раздел «Контрагенты»", level=2)
    add_screenshot_note(
        doc,
        caption="Справочник контрагентов",
        what_to_capture="Скриншот страницы «Контрагенты» с ИНН, наименованиями, адресами и телефонами.",
    )

    add_heading(doc, "8.6. Раздел «Бюджет»", level=2)
    add_para(
        doc,
        "Раздел «Бюджет» отражает плановые и фактические суммы расходов по "
        "отделам и статьям в разрезе года и квартала. Для каждой записи "
        "вычисляется отклонение «План − Факт».",
    )
    add_screenshot_note(
        doc,
        caption="Таблица бюджета",
        what_to_capture="Скриншот страницы «Бюджет» с таблицей план/факт по отделам и статьям.",
    )

    add_heading(doc, "8.7. Раздел «Документы»", level=2)
    add_para(
        doc,
        "Раздел «Документы» содержит первичные документы со связями: отдел, "
        "статья бюджета, тип документа, контрагент, ответственный сотрудник. "
        "Поддерживается создание нового документа, редактирование и удаление.",
    )
    add_screenshot_note(
        doc,
        caption="Список документов",
        what_to_capture="Скриншот страницы «Документы» с таблицей и применёнными фильтрами/сортировкой.",
    )
    add_screenshot_note(
        doc,
        caption="Удаление документа — окно подтверждения",
        what_to_capture="Нажмите «Удалить» у любой записи и сделайте скриншот окна подтверждения удаления.",
    )

    # ── OLAP ────────────────────────────────────────────────────────────────
    add_heading(doc, "9. OLAP-анализ данных", level=1)
    add_para(
        doc,
        "На странице «OLAP» реализованы основные операции многомерного анализа "
        "данных над бюджетом и документами. Измерения: отдел, статья бюджета, "
        "контрагент, год, квартал. Меры: плановая сумма, фактическая сумма, "
        "отклонение и количество документов.",
    )

    add_heading(doc, "9.1. Roll-up (агрегация)", level=2)
    add_para(
        doc,
        "Операция roll-up выполняет агрегацию мер по выбранному измерению. "
        "Реализованы три варианта: по отделам, по статьям бюджета, по "
        "кварталам. Запрос строится с помощью GROUP BY и агрегирующих "
        "функций SUM и COUNT.",
    )
    add_screenshot_note(
        doc,
        caption="Roll-up по отделам",
        what_to_capture="Раздел OLAP → вкладка «Roll-up по отделам». Скриншот таблицы (отдел, план, факт, отклонение, количество документов).",
    )
    add_screenshot_note(
        doc,
        caption="Roll-up по статьям бюджета",
        what_to_capture="Раздел OLAP → вкладка «Roll-up по статьям». Скриншот таблицы со статьями и их суммами.",
    )
    add_screenshot_note(
        doc,
        caption="Roll-up по кварталам",
        what_to_capture="Раздел OLAP → вкладка «Roll-up по кварталам». Скриншот с разбивкой по годам и кварталам.",
    )

    add_heading(doc, "9.2. Slice (срез)", level=2)
    add_para(
        doc,
        "Операция slice фиксирует одно измерение, возвращая срез "
        "многомерного куба. Реализованы срезы по конкретному отделу и по "
        "конкретному контрагенту: для выбранного значения выводятся все "
        "связанные бюджетные записи и документы.",
    )
    add_screenshot_note(
        doc,
        caption="Slice по отделу",
        what_to_capture="В разделе OLAP выберите «Срез по отделу», укажите конкретный отдел и сделайте скриншот результата (карточка отдела, его бюджет и документы).",
    )
    add_screenshot_note(
        doc,
        caption="Slice по контрагенту",
        what_to_capture="В разделе OLAP выберите «Срез по контрагенту», укажите конкретного контрагента и сделайте скриншот результата.",
    )

    add_heading(doc, "9.3. Dice (фильтрация по нескольким измерениям)", level=2)
    add_para(
        doc,
        "Операция dice позволяет задать фильтр сразу по нескольким измерениям "
        "(отдел, статья, год, квартал, контрагент). В SQL это реализовано "
        "через динамическое формирование WHERE-условия с параметрами.",
    )
    add_screenshot_note(
        doc,
        caption="Dice — применение нескольких фильтров",
        what_to_capture="В разделе OLAP → «Dice» задайте одновременно несколько фильтров (например, отдел + год + квартал) и сделайте скриншот окна с выбранными фильтрами и результирующей таблицей.",
    )

    add_heading(doc, "9.4. Drill-down (детализация)", level=2)
    add_para(
        doc,
        "Операция drill-down предоставляет детализированную информацию по "
        "выбранному отделу: сводная статистика, список сотрудников отдела, "
        "бюджет и документы. Переход от агрегата к деталям выполняется в один клик.",
    )
    add_screenshot_note(
        doc,
        caption="Drill-down по отделу",
        what_to_capture="В OLAP-разделе откройте детализацию по отделу. Скриншот должен включать сводную панель (total_plan, total_fact, deviation, количество сотрудников/документов) и таблицы сотрудников и документов.",
    )

    add_heading(doc, "9.5. Cross-tab (кросс-таблица «Отдел × Статья»)", level=2)
    add_para(
        doc,
        "Кросс-таблица отображает суммарные плановые и фактические значения "
        "в разрезе двух измерений одновременно — отдел и статья бюджета. "
        "Позволяет быстро выявить, на какие статьи и в каких отделах "
        "приходится основная доля расходов.",
    )
    add_screenshot_note(
        doc,
        caption="Кросс-таблица «Отдел × Статья»",
        what_to_capture="В разделе OLAP → «Cross-tab» сделайте скриншот сводной таблицы «Отдел × Статья» со значениями плана и факта.",
    )

    # ── Визуализация ────────────────────────────────────────────────────────
    add_heading(doc, "10. Визуализация данных (диаграммы)", level=1)
    add_para(
        doc,
        "Для наглядного представления результатов OLAP-анализа в приложении "
        "реализованы диаграммы (компонент OlapChart). Они позволяют сравнить "
        "план и факт по отделам и статьям, оценить структуру расходов и "
        "динамику по кварталам.",
    )
    add_screenshot_note(
        doc,
        caption="Диаграмма «План vs Факт по отделам»",
        what_to_capture="В OLAP-разделе переключитесь в режим диаграммы для roll-up по отделам и сделайте скриншот столбчатой (или круговой) диаграммы.",
    )
    add_screenshot_note(
        doc,
        caption="Диаграмма «Динамика по кварталам»",
        what_to_capture="Скриншот диаграммы по кварталам, где видна динамика плана/факта во времени.",
    )

    # ── Экспорт ─────────────────────────────────────────────────────────────
    add_heading(doc, "11. Экспорт отчётов", level=1)
    add_para(
        doc,
        "Для всех аналитических отчётов предусмотрен экспорт в форматы CSV "
        "и XLSX. Выгрузка выполняется на стороне сервера (модуль export.py): "
        "SQL-запрос формирует строки с русскоязычными заголовками столбцов, "
        "затем они сериализуются в нужный формат (openpyxl для XLSX, "
        "встроенная библиотека csv для CSV) и отдаются клиенту через "
        "StreamingResponse с заголовком Content-Disposition.",
    )
    add_screenshot_note(
        doc,
        caption="Кнопки экспорта отчёта",
        what_to_capture="В разделе OLAP сделайте скриншот панели с кнопками «Экспорт CSV» / «Экспорт XLSX».",
    )
    add_screenshot_note(
        doc,
        caption="Выгруженный XLSX-файл",
        what_to_capture="Откройте выгруженный файл (например, rollup-by-dept.xlsx) в MS Excel и сделайте скриншот содержимого со всеми столбцами и заголовками на русском языке.",
    )
    add_screenshot_note(
        doc,
        caption="Выгруженный CSV-файл",
        what_to_capture="Откройте выгруженный CSV в текстовом редакторе (или Excel) и сделайте скриншот содержимого.",
    )

    # ── Тестирование ────────────────────────────────────────────────────────
    add_heading(doc, "12. Тестирование и проверка корректности", level=1)
    add_para(
        doc,
        "Для наполнения БД использован сценарий seed_data.sql, добавляющий "
        "репрезентативный набор данных: 15 должностей, 8 типов документов, "
        "12 статей бюджета, 10 контрагентов, отделы, сотрудников, бюджетные "
        "записи и документы. Проверка корректности выполнена на следующих "
        "сценариях:",
    )
    add_bullet(doc, "создание, изменение и удаление записей во всех разделах;")
    add_bullet(doc, "соблюдение ограничений целостности — при попытке удаления справочника, на который ссылаются другие записи, сервер возвращает ошибку;")
    add_bullet(doc, "корректность агрегации: суммы roll-up совпадают с суммами детальных срезов;")
    add_bullet(doc, "корректность фильтрации dice при произвольной комбинации параметров;")
    add_bullet(doc, "корректность выгрузки: содержимое CSV/XLSX совпадает с отображаемой в интерфейсе таблицей.")

    add_screenshot_note(
        doc,
        caption="Пример валидации — ошибка при некорректных данных",
        what_to_capture="Попробуйте создать запись с некорректными данными (пустое обязательное поле, дубликат PK) и сделайте скриншот сообщения об ошибке от API или в интерфейсе.",
    )
    add_screenshot_note(
        doc,
        caption="Результат SQL-запроса в pgAdmin",
        what_to_capture="Выполните в pgAdmin запрос вида SELECT COUNT(*) FROM employee; или произвольный JOIN-запрос и сделайте скриншот результата, чтобы подтвердить соответствие данных в БД и в UI.",
    )

    # ── Выводы ──────────────────────────────────────────────────────────────
    add_heading(doc, "13. Выводы", level=1)
    add_para(
        doc,
        "В ходе выполнения лабораторной работы разработано полнофункциональное "
        "клиент-серверное приложение для работы с реляционной базой данных "
        "PostgreSQL. Реализованы все поставленные задачи:",
    )
    add_bullet(doc, "спроектирована схема БД из восьми взаимосвязанных таблиц, отражающая предметную область документооборота и бюджетирования;")
    add_bullet(doc, "реализован REST API на FastAPI с полным набором CRUD-операций для каждой сущности;")
    add_bullet(doc, "разработан клиент на Vue 3 с удобным интерфейсом: таблицы с поиском и сортировкой, модальные формы ввода, навигация;")
    add_bullet(doc, "реализованы OLAP-операции (roll-up, slice, dice, drill-down, cross-tab) с использованием агрегирующих SQL-запросов;")
    add_bullet(doc, "добавлена визуализация аналитических данных в виде диаграмм;")
    add_bullet(doc, "реализован экспорт отчётов в форматы CSV и XLSX.")
    add_para(
        doc,
        "Таким образом, цель лабораторной работы достигнута: приобретены "
        "практические навыки проектирования реляционных БД, разработки "
        "серверной и клиентской частей приложения, написания аналитических "
        "SQL-запросов для OLAP-анализа и организации экспорта данных.",
    )

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
